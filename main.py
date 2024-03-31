import streamlit as st
from streamlit_extras.stylable_container import stylable_container
from streamlit_extras.add_vertical_space import add_vertical_space
from deta import Deta
from io import BytesIO
import random

# PPT 라이브러리
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# 시간 정보 라이브러리
from datetime import datetime
import pytz

# g4f 사용하여 Grammar Check
from g4f.client import Client
from g4f.Provider import GptTalkRu

st.set_page_config(page_title='Write Your Essay', page_icon='✏️')

# Time Zone을 서울로 설정
seoul = pytz.timezone('Asia/Seoul')

# style.css 파일 열고 적용
with open('style.css', encoding='UTF-8') as f:
    st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

# 사용자의 고유 아이디
if 'key' not in st.session_state:
    st.session_state['key'] = random.randint(1,100000000000000)


st.title('Write Your Essay')

# 성품 선택
CharacterTraits = ['Affectionate', 'Appreciative', 'Attentive', 'Available', 'Blessed', 'Cheerful', 'Committed', 'Compassionate', 'Concerned', 'Confident', 'Considerate', 'Consistent', 'Content', 'Cooperative', 'Courageous', 'Courteous', 'Creative', 'Decisive', 'Deferent', 'Dependable', 'Determined', 'Diligent', 'Discerning', 'Discreet', 'Efficient', 'Equitable', 'Fair', 'Faithful', 'Fearless', 'Flexible', 'Forgiving', 'Friendly', 'Generous', 'Gentle', 'Godly', 'Goodly', 'Gracious', 'Grateful', 'Happy', 'Holy', 'Honest', 'Humble', 'Integrity', 'Joyful', 'Just', 'Kind', 'Knowledgeable', 'Longsuffering', 'Loving', 'Loyal', 'Meek', 'Merciful', 'Modest', 'Obedient', 'Observant', 'Optimistic', 'Orderly', 'Patient', 'Peaceful', 'Perseverant', 'Persuasive', 'Prepared', 'Prudent', 'Punctual', 'Pure', 'Purposeful', 'Ready', 'Rejoiceful', 'Resourceful', 'Respectful', 'Responsible', 'Reverent', 'Righteous', 'Secure', 'Self-Controlled', 'Sincere', 'Steadfast', 'Submissive', 'Tactful', 'Temperate', 'Thorough', 'Thrifty', 'Tolerant', 'Trustworthy', 'Truthful', 'Understanding', 'Virtuous', 'Wise', 'Zealous']
col1, col2 = st.columns(2)
with col1: 
    if 'random_topic' not in st.session_state:
        st.session_state.random_topic = 0
    st.session_state.topic = st.selectbox('성품을 골라주세요', CharacterTraits, index=st.session_state.random_topic)
with col2: 
    '####'
    topic2 = st.button('랜덤 성품 뽑기')

# 랜덤 성품
if topic2:
    st.session_state.random_topic = random.randint(0, len(CharacterTraits)-1)
    st.session_state.topic = st.session_state.random_topic
    st.rerun()

# 이름 작성
st.session_state.name = st.text_input('이름을 작성해주세요')

st.divider()

if 'words' not in st.session_state:
    st.session_state.words = 0

st.session_state.content = st.text_area(f'에세이를 작성해주세요', placeholder=f'{st.session_state.topic}에 대하여 영어로 에세이를 작성해보세요', height=300)

col1, col2, col3 = st.columns(3)
with col1.popover("작성법", use_container_width=1):
    st.subheader('에세이 작성하는 법')
    '''
    선택한 성품이 무엇인지 쓰고,  
    해당 성품에 대한 생각과 느낀점을  
    자유롭게 150단어 이상 영어로 작성해보세요!😀'''

if col2.button('단어 갯수', use_container_width=1):
    st.session_state.words = len(st.session_state.content.split())
    with col3:
        with stylable_container(
            key="container_with_border",
            css_styles="""
                {
                    padding-top: 0.4rem;
                    padding-bottom: 0.3rem;
                    border: 1px solid rgba(49, 51, 63, 0.2);
                    border-radius: 0.5rem;
                    text-align: center;
                }
                """,
        ):
            st.write(f"현재 **{st.session_state.words}**단어")

if st.button('Check Grammar'):
    client = Client(provider=GptTalkRu)
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "From now on, you are a grammar checker. Check some errors in the sentence and label the incorrect word with :red[]."},
            {"role": "user", "content": "i is good at workin computer?"},
            {"role": "assistant", "content": ":red[i] :red[is] good at :red[workin] computer:red[?]"},
            {"role": "user", "content": st.session_state.content},
            ],
    )
    st.write(response.choices[0].message.content)
    


with stylable_container(
    key="green_button",
    css_styles="""
        button {
            background: linear-gradient(45deg, #c5d1f6, #ffc7ca);
            color: black;
            border-radius: 20px;
        }
        """,
):
    submit = st.button("제이크 선생님에게 제출")

# submit 버튼이 눌렸을 때 실행
if submit:
    # 혹시 모르니깐 단어 갯수 확인 ㄱㄱ
    st.session_state.words = len(st.session_state.content.split())
    # 사용자가 모든 빈칸들을 채웠는지 확인
    if st.session_state.name == '':
        st.error('이름을 작성해주세요')
    elif st.session_state.words < 150:
        st.error('150단어 이상 작성해주세요')
    else:
        # 문서 파일 생성
        doc = Document()
        doc.add_heading(st.session_state.topic, level=0)
        today = datetime.now(tz=seoul).date()
        doc.add_heading(f'{st.session_state.name}, {today}', level=1)
        doc.add_paragraph(st.session_state.content)

        # 메모리 내에서 파일을 BytesIO 객체로 저장
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Deta Base에 파일 업로드
        def upload_file(file_stream):
            DETA_KEY = 'c0ki5D3avML_gSssDuj33rfuzLDrjwL1gc42oQkbgsHj'
            deta = Deta(DETA_KEY)
            db = deta.Drive("Write_Your_Essay")

            # 파일 이름 설정
            file_name = f'{today}_{st.session_state.name}_{st.session_state.topic}.docx'
            
            # 파일 업로드
            db.put(file_name, file_stream)

        # 업로드 실행
        with st.spinner('업로드 중...'):
            upload_file(file_stream)
        st.success('제이크 선생님에게 제출되었습니다!')
